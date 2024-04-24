import requests
from PIL import Image
from io import BytesIO
from googleapiclient.discovery import build
from docx import Document
from docx.shared import Inches
from PIL import UnidentifiedImageError

def ambil_hasil_pencarian(kunci_api, kata_kunci):
    service = build("customsearch", "v1", developerKey=kunci_api)
    hasil = service.cse().list(q=kata_kunci, cx='50b1bcd4363224c93').execute()
    items = hasil.get('items', [])
    
    hasil_pencarian = []
    for item in items:
        judul = item.get('title', '')
        deskripsi = item.get('snippet', '')
        gambar_url = None
        if 'pagemap' in item and 'cse_image' in item['pagemap']:
            gambar_url = item['pagemap']['cse_image'][0]['src']
        hasil_pencarian.append({'judul': judul, 'deskripsi': deskripsi, 'gambar_url': gambar_url})
    
    return hasil_pencarian

def simpan_gambar(url, nama_file):
    try:
        response = requests.get(url)
        if response.status_code == 200:
            image = Image.open(BytesIO(response.content))
            image.save(nama_file)
    except UnidentifiedImageError as e:
        print(f"Gagal menyimpan gambar: {e}")
    except Exception as e:
        print(f"Terjadi kesalahan saat menyimpan gambar: {e}")

def buat_dokumen(judul, hasil_pencarian, nama_file):
    doc = Document()
    doc.add_heading(judul, level=1)
    
    for index, item in enumerate(hasil_pencarian, start=1):
        try:
            doc.add_paragraph(f"{index}. {item['deskripsi']}")
            if item['gambar_url']:
                gambar_nama_file = f"gambar_{judul}_{index}.png"
                simpan_gambar(item['gambar_url'], gambar_nama_file)
                doc.add_picture(gambar_nama_file, width=Inches(2))
        except UnicodeEncodeError:
            deskripsi = item['deskripsi'].encode('latin-1', 'ignore').decode('latin-1')
            doc.add_paragraph(f"{index}. {deskripsi}")
        except Exception as e:
            print(f"Terjadi kesalahan saat menambahkan gambar: {e}")
    
    doc.save(nama_file)

# Gunakan fungsi untuk mendapatkan hasil pencarian
kunci_api_google = 'AIzaSyBJWD7KUDyauC1LPW8uCA-w472IIk76eRc'
kode = '402'
hasil_pencarian = ambil_hasil_pencarian(kunci_api_google, judul_pencarian)

# Buat 5 dokumen teks dengan hasil pencarian
for i in range(5):
    nama_file = f"hasil_pencarian_{judul_pencarian}_{i+1}.docx"
    buat_dokumen(judul_pencarian, hasil_pencarian, nama_file)