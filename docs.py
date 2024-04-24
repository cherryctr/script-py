from docx import Document
from docx.shared import Inches
import requests
from bs4 import BeautifulSoup
import os

def create_document(question, answer_text, image_urls, file_name):
    doc = Document()

    # Add question
    doc.add_heading('Pertanyaan:', level=1)
    doc.add_paragraph(question)

    # Add answer text
    doc.add_heading('Jawaban:', level=1)
    doc.add_paragraph(answer_text)

    # Add images
    doc.add_heading('Langkah-langkah:', level=1)
    for idx, url in enumerate(image_urls):
        img_data = requests.get(url).content
        img_path = f"image_{idx+1}.png"
        with open(img_path, 'wb') as img_file:
            img_file.write(img_data)
        doc.add_picture(img_path, width=Inches(5))

    # Save document
    doc.save(file_name)
    print(f"Dokumen '{file_name}' berhasil dibuat.")

def search_google(query):
    url = f"https://www.google.com/search?q={'+'.join(query.split())}"
    headers = {'User-Agent': 'Mozilla/5.0'}
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.text, 'html.parser')
    results = soup.find_all('div', class_='BNeawe vvjwJb AP7Wnd')
    if results:
        return results[0].get_text()
    else:
        return "Maaf, tidak ada jawaban yang ditemukan."

def get_image_urls(query):
    url = f"https://www.google.com/search?q={'+'.join(query.split())}&tbm=isch"
    headers = {'User-Agent': 'Mozilla/5.0'}
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.text, 'html.parser')
    image_tags = soup.find_all('img', class_='t0fcAb')
    return [img['src'] for img in image_tags]

def main():
    question = input("Masukkan pertanyaan: ")
    answer_text = search_google(question)
    image_urls = get_image_urls(question)
    create_document(question, answer_text, image_urls, "Jawaban_Google_Ads.docx")

if __name__ == "__main__":
    main()
