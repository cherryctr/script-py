import os
import pandas as pd
from selenium import webdriver
import requests
import io
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx2pdf import convert

# Set the path to the Chrome WebDriver executable
PATH = "C:\\Users\\Admin\\Documents\\pysc\\packages\\chromedriver.exe"

# Initialize Chrome driver
wd = webdriver.Chrome(PATH)

# Sample image URL
image_url = "https://id.pinterest.com/search/pins/?q=eccomerce&rs=typed"

def download_image(download_path, url, file_name):
    try:
        image_content = requests.get(url).content
        image_file = io.BytesIO(image_content)
        image = Image.open(image_file)

        file_path = os.path.join(download_path, file_name)

        with open(file_path, "wb") as f:
            image.save(f, "JPEG")

        print("SUCCESS")

        # Export to a Word document
        doc = Document()
        doc.add_heading(f'Jawaban {i}', 0)  # Modify the heading to use the loop index
        doc.add_picture(file_path, width=Inches(4))
        # doc.add_paragraph("Image source: " + url)
        doc.save(file_path)
        print(f"Exported to Word document: {file_path}")

        # Convert Word document to PDF
        convert(file_path)

    except Exception as e:
        print(f"Error downloading image: {e}")
        print(f"Failed to download image from URL: {url}")

# Create the folder "HASIL JAWABAN" if it doesn't exist
output_folder = "HASIL JAWABAN"
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# Load the list of names from an Excel file
names_df = pd.read_excel("C:\\Users\\Admin\\Documents\\pysc\\excels\\DAFTAR_NAMA_TES.xlsx")

# Initialize a dictionary to store the count for each name
name_count = {}

# Iterate through each name
for index, row in names_df.iterrows():
    name = row['Nama']
    # Check if the name is in the dictionary
    if name not in name_count:
        name_count[name] = 1  # Initialize the count to 1 for a new name
    else:
        name_count[name] += 1  # Increment the count for an existing name
    # Get the count for the current name
    count = name_count[name]
    for i in range(1, 6):  # Create 5 documents for each name
        file_name = f"JAWABAN_{name.upper()}_{i}.docx"
        download_image(output_folder, image_url, file_name)
